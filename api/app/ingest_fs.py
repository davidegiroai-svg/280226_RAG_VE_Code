"""
Filesystem ingest worker (M0)
- Creates/gets KB by namespace
- Inserts documents with dedup (kb_id, content_hash)
- Splits text into chunks (size 1200, overlap 200)
- Inserts chunks with metadata including source_path
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
from pathlib import Path
from typing import Iterable, Tuple, List

import logging

import psycopg2
from psycopg2.extras import Json

from app.embedding import embed_texts, EmbeddingError
from app.entity_extractor import extract_and_save as _graph_extract_and_save

logger = logging.getLogger(__name__)


def _env(name: str, default: str) -> str:
    v = os.getenv(name)
    return v if v not in (None, "") else default


def get_conn():
    host = _env("DB_HOST", "db")
    port = int(_env("DB_PORT", "5432"))
    user = _env("POSTGRES_USER", "rag")
    pwd = _env("POSTGRES_PASSWORD", "rag_password_change_me")
    db = _env("POSTGRES_DB", "rag")

    return psycopg2.connect(host=host, port=port, user=user, password=pwd, dbname=db)


def sha256_text(text: str) -> str:
    h = hashlib.sha256()
    h.update(text.encode("utf-8", errors="ignore"))
    return h.hexdigest()


MAX_PREFIX_LEN = 150


def build_context_prefix(
    kb_namespace: str,
    titolo: str,
    file_type: str,
    tipo_documento: str = "",
    targets: list | None = None,
) -> str:
    """Costruisce prefisso contestuale per embedding (zero LLM cost).

    Formato: "[kb | titolo | tipo | tipo_doc | target1 target2] "
    Troncato a MAX_PREFIX_LEN per lasciare spazio al testo (MAX_CHARS=2000).
    """
    parts = [p for p in [kb_namespace, titolo, file_type] if p]
    if tipo_documento:
        parts.append(tipo_documento)
    if targets:
        parts.append(" ".join(targets[:3]))
    prefix = "[" + " | ".join(parts) + "] "
    return prefix[:MAX_PREFIX_LEN]


def read_text_file(p: Path) -> str:
    """Read text file with robust encoding handling.

    Tries utf-8-sig first (removes BOM), falls back to utf-8 with errors=ignore.
    Also includes heuristic repair for mojibake (e.g. "Ã¨" -> "è").
    """
    try:
        # Try utf-8-sig first (auto-removes BOM)
        content = p.read_text(encoding="utf-8-sig", errors="ignore")
    except Exception:
        # Fallback to utf-8 with errors=ignore
        content = p.read_text(encoding="utf-8", errors="ignore")
    # Explicit BOM removal as safety net
    content = content.replace("\ufeff", "")
    # Heuristic mojibake repair: detect typical patterns and attempt round-trip
    # Only activate if we see common mojibake indicators
    if "Ã" in content or "Â" in content:
        try:
            # Try latin1 -> utf-8 round trip to fix mojibake
            repaired = content.encode("latin-1", errors="ignore").decode("utf-8", errors="ignore")
            # Only accept if it looks better (fewer mojibake patterns)
            if "Ã" not in repaired and "Â" not in repaired:
                content = repaired
        except Exception:
            # Fail silently if repair doesn't work
            pass
    return content


def read_sidecar_meta(doc_path: Path) -> dict:
    """Legge il file .meta.json sidecar accanto al documento se esiste.

    Cerca <nome_documento>.meta.json nella stessa directory del file.
    Ritorna il dict con i metadati, oppure {} se assente o malformato.
    """
    meta_path = doc_path.with_suffix(".meta.json")
    if not meta_path.exists():
        return {}
    try:
        return json.loads(meta_path.read_text(encoding="utf-8"))
    except Exception:
        return {}


def read_docx_file(p: Path) -> str:
    """Estrae testo da file DOCX: paragrafi + celle tabelle, separati da newline."""
    import docx
    doc = docx.Document(str(p))
    parts = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def chunk_text_legacy(text: str, size: int = 3000, overlap: int = 500) -> Iterable[Tuple[int, str]]:
    text = text.strip()
    if not text:
        return
    if overlap >= size:
        overlap = max(0, size // 5)

    start = 0
    idx = 0
    n = len(text)
    step = max(1, size - overlap)

    while start < n:
        end = min(start + size, n)
        yield idx, text[start:end]
        idx += 1
        if end >= n:
            break
        start += step


def chunk_text_semantic(
    text: str,
    target_size: int = 1500,
    max_size: int = 2000,
    overlap_sentences: int = 2,
) -> List[Tuple[int, str]]:
    """Chunking semantico: spezza su fine frase, rispetta target_size/max_size.

    Parametri:
        target_size: dimensione target del chunk in caratteri
        max_size: dimensione massima assoluta prima di forzare il taglio
        overlap_sentences: quante frasi dell'ultimo chunk ripetere nel successivo
    """
    text = text.strip()
    if not text:
        return []

    # Split su fine frase (. ! ? seguiti da spazio o newline)
    sentences = re.split(r'(?<=[.!?])\s+', text)
    sentences = [s for s in sentences if s.strip()]
    if not sentences:
        return []

    chunks: List[Tuple[int, str]] = []
    current: List[str] = []
    current_len = 0
    idx = 0

    for sent in sentences:
        sent_len = len(sent)
        if current and current_len + sent_len > max_size:
            chunks.append((idx, " ".join(current)))
            idx += 1
            # Overlap: mantieni ultime N frasi come contesto per il chunk successivo
            if overlap_sentences and len(current) >= overlap_sentences:
                current = current[-overlap_sentences:]
                current_len = sum(len(s) for s in current)
            else:
                current = []
                current_len = 0
        current.append(sent)
        current_len += sent_len

    if current:
        chunks.append((idx, " ".join(current)))

    return chunks


# Selettore chunk mode via env CHUNK_MODE=semantic|legacy (default: semantic)
def chunk_text(text: str, size: int = 3000, overlap: int = 500) -> Iterable[Tuple[int, str]]:
    """Wrapper: usa chunk_text_semantic o chunk_text_legacy in base a CHUNK_MODE env."""
    mode = _env("CHUNK_MODE", "semantic")
    if mode == "legacy":
        yield from chunk_text_legacy(text, size=size, overlap=overlap)
    else:
        yield from chunk_text_semantic(text)


def ensure_kb(cur, namespace: str) -> str:
    cur.execute("SELECT id::text FROM knowledge_base WHERE namespace = %s", (namespace,))
    row = cur.fetchone()
    if row:
        return row[0]

    cur.execute(
        "INSERT INTO knowledge_base (namespace, nome, descrizione) VALUES (%s, %s, %s) RETURNING id::text",
        (namespace, namespace, f"KB auto-created for namespace '{namespace}'"),
    )
    return cur.fetchone()[0]


def upsert_document(cur, kb_id: str, source_uri: str, titolo: str, content_hash: str):
    cur.execute(
        """
        INSERT INTO documents (kb_id, source_uri, titolo, content_hash)
        VALUES (%s, %s, %s, %s)
        ON CONFLICT (kb_id, content_hash) DO NOTHING
        RETURNING id::text
        """,
        (kb_id, source_uri, titolo, content_hash),
    )
    row = cur.fetchone()
    if row:
        return row[0], True

    cur.execute("SELECT id::text FROM documents WHERE kb_id=%s AND content_hash=%s", (kb_id, content_hash))
    row2 = cur.fetchone()
    if not row2:
        raise RuntimeError("Cannot retrieve existing document after conflict")
    return row2[0], False


def vector_to_str(vec: List[float]) -> str:
    """Convert Python list to PostgreSQL vector string format."""
    return "[" + ",".join(str(x) for x in vec) + "]"


def insert_chunks(
    cur,
    kb_id: str,
    kb_namespace: str,
    doc_id: str,
    source_path: str,
    file_name: str,
    text: str,
    *,
    file_path: Path = None,
    titolo: str = "",
) -> int:
    # Leggi metadati sidecar se disponibili
    sidecar = read_sidecar_meta(file_path) if file_path is not None else {}

    # Deriva file_type dall'estensione (es. ".pdf" → "pdf")
    _file_ext = (file_path.suffix.lower().lstrip(".") if file_path is not None
                 else Path(file_name).suffix.lower().lstrip("."))

    # Prefisso contestuale per embedding (testo in DB resta invariato)
    prefix = build_context_prefix(
        kb_namespace, titolo, _file_ext,
        tipo_documento=sidecar.get("tipo_documento", ""),
        targets=sidecar.get("targets"),
    )

    # Branch PDF: usa read_pdf_chunks con page_start/page_end come colonne dedicate
    if file_path is not None and file_path.suffix.lower() == ".pdf":
        page_chunks = read_pdf_chunks(file_path)
        valid_chunks = [(i, pc) for i, pc in enumerate(page_chunks) if pc["testo"].strip()]
        if not valid_chunks:
            return 0

        chunk_texts_list = [prefix + pc["testo"] for _, pc in valid_chunks]
        try:
            embeddings, embedding_model, embedding_dim = embed_texts(chunk_texts_list)
        except EmbeddingError as e:
            raise RuntimeError(f"Embedding failed for PDF '{file_name}': {e}")

        inserted = 0
        for (chunk_index, pc), embedding in zip(valid_chunks, embeddings):
            meta = {
                "source_path": source_path,
                "file_name": file_name,
                "file_type": _file_ext,
                "chunk_index": chunk_index,
                "page_start": pc["page_start"],
                "page_end": pc["page_end"],
                **sidecar,
            }
            cur.execute(
                """
                INSERT INTO chunks (
                    document_id, kb_id, kb_namespace, chunk_index, testo,
                    metadata, embedding, embedding_model, embedding_dim,
                    page_start, page_end, section_title, doc_title
                )
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                """,
                (
                    doc_id, kb_id, kb_namespace, chunk_index, pc["testo"],
                    Json(meta), vector_to_str(embedding), embedding_model, embedding_dim,
                    pc["page_start"], pc["page_end"], pc["section_title"], titolo,
                ),
            )
            inserted += 1
        return inserted

    # Comportamento originale per TXT/MD/CSV/JSON
    chunks_data = []
    for chunk_index, chunk in chunk_text(text, 3000, 500):
        chunks_data.append((chunk_index, chunk))

    if not chunks_data:
        return 0

    chunk_texts_list = [prefix + c[1] for c in chunks_data]
    try:
        embeddings, embedding_model, embedding_dim = embed_texts(chunk_texts_list)
    except EmbeddingError as e:
        raise RuntimeError(f"Embedding failed for document '{file_name}': {e}")

    inserted = 0
    for (chunk_index, chunk), embedding in zip(chunks_data, embeddings):
        meta = {"source_path": source_path, "file_name": file_name, "file_type": _file_ext, "chunk_index": chunk_index, **sidecar}
        cur.execute(
            """
            INSERT INTO chunks (document_id, kb_id, kb_namespace, chunk_index, testo, metadata, embedding, embedding_model, embedding_dim, doc_title)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """,
            (doc_id, kb_id, kb_namespace, chunk_index, chunk, Json(meta), vector_to_str(embedding), embedding_model, embedding_dim, titolo),
        )
        inserted += 1
    return inserted


def update_ingest_status(cur, doc_id: str, status: str) -> None:
    """Aggiorna lo stato di ingest di un documento.

    Args:
        cur:    cursore psycopg2 aperto.
        doc_id: UUID del documento.
        status: nuovo stato ('pending', 'processing', 'done', 'error').
    """
    cur.execute(
        "UPDATE documents SET ingest_status=%s, updated_at=now() WHERE id=%s",
        (status, doc_id),
    )


def _auto_classify(file_path: Path, text_snippet: str) -> None:
    """Auto-classificazione: se no sidecar, estrai metadati via LLM (best-effort)."""
    if file_path is None:
        return
    if os.getenv("AUTO_CLASSIFY_ENABLED", "false").lower() != "true":
        return
    existing_sidecar = read_sidecar_meta(file_path)
    if existing_sidecar:
        return
    try:
        from app.metadata_extractor import extract_metadata_for_file, save_sidecar_meta
        model = os.getenv("OLLAMA_LLM_MODEL", "qwen3-next-cloud:latest")
        timeout = int(os.getenv("METADATA_EXTRACT_TIMEOUT", "120"))
        extracted = extract_metadata_for_file(file_path, model, text_snippet, timeout=timeout)
        if extracted.get("tipo_documento"):  # non salvare fallback vuoto
            save_sidecar_meta(file_path, extracted)
            logger.info("Auto-classificazione completata per %s", file_path.name)
    except Exception:
        pass  # best-effort, non blocca ingest


def ingest_single_file(file_path: Path, kb_namespace: str) -> dict:
    """Ingestisce un singolo file nel DB. Usato dal watcher per auto-ingest.

    Ciclo di vita dello stato: pending (inserimento) → processing → done/error.

    Args:
        file_path:    Path assoluta al file da ingestire.
        kb_namespace: Namespace della KB di destinazione.

    Returns:
        dict con chiavi: status, doc_id, is_new, chunks_inserted.

    Raises:
        RuntimeError: se l'ingest fallisce.
    """
    ext = file_path.suffix.lower()
    supported = {".txt", ".md", ".csv", ".json", ".pdf", ".docx"}
    if ext not in supported:
        return {"status": "skipped", "reason": "estensione non supportata"}

    conn = get_conn()
    conn.autocommit = False

    try:
        with conn.cursor() as cur:
            kb_id = ensure_kb(cur, kb_namespace)

            is_pdf = ext == ".pdf"
            is_docx = ext == ".docx"
            if is_pdf:
                raw_bytes = file_path.read_bytes()
                content_hash = hashlib.sha256(raw_bytes).hexdigest()
                text = ""
            elif is_docx:
                text = read_docx_file(file_path)
                if not text.strip():
                    conn.rollback()
                    return {"status": "skipped", "reason": "file vuoto"}
                content_hash = sha256_text(text)
            else:
                text = read_text_file(file_path)
                if not text.strip():
                    conn.rollback()
                    return {"status": "skipped", "reason": "file vuoto"}
                content_hash = sha256_text(text)

            source_path = file_path.as_posix()
            doc_id, is_new = upsert_document(cur, kb_id, source_path, file_path.name, content_hash)

            if not is_new:
                conn.rollback()
                return {"doc_id": doc_id, "is_new": False, "chunks_inserted": 0, "status": "existing"}

            # Transizione stato: processing
            update_ingest_status(cur, doc_id, "processing")

            # Auto-classificazione: se no sidecar, estrai metadati via LLM
            _auto_classify(file_path, text)

            chunks_inserted = insert_chunks(
                cur, kb_id, kb_namespace, doc_id, source_path, file_path.name, text,
                file_path=file_path, titolo=file_path.name,
            )

            # Transizione stato: done
            update_ingest_status(cur, doc_id, "done")

            # M7 GraphRAG: estrazione entità (best-effort, non blocca mai l'ingest)
            _extraction_text = text
            if is_pdf:
                try:
                    _pdf_pages = read_pdf_chunks(file_path)
                    _extraction_text = " ".join(pc["testo"] for pc in _pdf_pages)
                except Exception:
                    _extraction_text = ""
            _graph_extract_and_save(cur, doc_id, kb_id, _extraction_text)

        conn.commit()
        return {"doc_id": doc_id, "is_new": True, "chunks_inserted": chunks_inserted, "status": "done"}

    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()


def read_pdf_chunks(p: Path) -> list:
    """Legge un PDF con pymupdf4llm e restituisce chunk sub-pagina.

    Ogni pagina viene ulteriormente suddivisa con chunk_text() (size=800, overlap=150)
    per creare embedding più focalizzati e migliorare la qualità del retrieval.
    Pagine corte (< 800 char) producono un unico chunk invariato.

    Ritorna lista di dict con chiavi: testo, page_start, page_end, section_title.
    """
    import pymupdf4llm
    pages = pymupdf4llm.to_markdown(str(p), page_chunks=True)
    result = []
    for page in pages:
        page_text = page["text"].strip()
        if not page_text:
            continue
        page_num = page["metadata"]["page"]
        for _, sub_text in chunk_text_legacy(page_text, size=800, overlap=150):
            if sub_text.strip():
                result.append({
                    "testo": sub_text,
                    "page_start": page_num,
                    "page_end": page_num,
                    "section_title": None,
                })
    return result


def list_files(root: Path):
    exts = {".txt", ".md", ".csv", ".json", ".pdf", ".docx"}
    for p in root.rglob("*"):
        # Esclude i file sidecar .meta.json che non devono essere ingestiti
        if p.name.endswith(".meta.json"):
            continue
        if p.is_file() and p.suffix.lower() in exts:
            yield p


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--kb", default="demo")
    ap.add_argument("--path", default="/data/inbox")
    args = ap.parse_args()

    kb_namespace = args.kb.strip()
    in_path = Path(args.path)

    if not kb_namespace:
        raise SystemExit("ERROR: --kb is empty")
    if not in_path.exists() or not in_path.is_dir():
        raise SystemExit(f"ERROR: path not found or not a dir: {in_path}")

    files = list(list_files(in_path))
    if not files:
        print(f"INFO: No supported files found in {in_path}")
        return

    files_read = 0
    docs_new = 0
    docs_skipped = 0
    chunks_inserted = 0

    # Ottieni kb_id in una transazione separata
    conn = get_conn()
    conn.autocommit = False
    try:
        with conn.cursor() as cur:
            kb_id = ensure_kb(cur, kb_namespace)
        conn.commit()
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()

    # Ingest file per file — commit separato per evitare OOM su transazioni grandi
    for fp in files:
        files_read += 1
        ext = fp.suffix.lower()
        is_pdf = ext == ".pdf"
        is_docx = ext == ".docx"

        if is_pdf:
            raw_bytes = fp.read_bytes()
            content_hash = hashlib.sha256(raw_bytes).hexdigest()
            text = ""
        elif is_docx:
            text = read_docx_file(fp)
            if not text.strip():
                continue
            content_hash = sha256_text(text)
        else:
            text = read_text_file(fp)
            if not text.strip():
                continue
            content_hash = sha256_text(text)

        source_path = fp.as_posix()
        titolo = fp.name

        conn = get_conn()
        conn.autocommit = False
        try:
            with conn.cursor() as cur:
                doc_id, inserted_new = upsert_document(cur, kb_id, source_path, titolo, content_hash)
                if not inserted_new:
                    docs_skipped += 1
                    conn.rollback()
                    conn.close()
                    continue

                docs_new += 1

                # Auto-classificazione: se no sidecar, estrai metadati via LLM
                _auto_classify(fp, text)

                n = insert_chunks(
                    cur, kb_id, kb_namespace, doc_id, source_path, fp.name, text,
                    file_path=fp, titolo=titolo,
                )
                chunks_inserted += n

                # M7 GraphRAG: estrazione entità (best-effort, non blocca mai l'ingest)
                _extraction_text = text
                if is_pdf:
                    try:
                        _pdf_pages = read_pdf_chunks(fp)
                        _extraction_text = " ".join(pc["testo"] for pc in _pdf_pages)
                    except Exception:
                        _extraction_text = ""
                _graph_extract_and_save(cur, doc_id, kb_id, _extraction_text)
            conn.commit()
            print(f"  OK {fp.name}: {n} chunks")
        except Exception as e:
            try:
                conn.rollback()
            except Exception:
                pass
            print(f"  ERR {fp.name}: {e}")
        finally:
            try:
                conn.close()
            except Exception:
                pass

    print("OK ingest completed")
    print(json.dumps({
        "kb": kb_namespace,
        "path": str(in_path),
        "files_found": len(files),
        "files_read": files_read,
        "documents_new": docs_new,
        "documents_skipped_existing": docs_skipped,
        "chunks_inserted": chunks_inserted
    }, indent=2))


if __name__ == "__main__":
    main()
