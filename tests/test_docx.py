"""tests/test_docx.py — Test DOCX parsing e ingestione (M5 Blocco 3)

Copre read_docx_file() e il comportamento di ingest_single_file() per edge case DOCX.
Usa python-docx (già in requirements.txt) per creare file DOCX programmaticamente.
Non richiede fixture su disco: tutto generato in tmp_path.
"""

import pytest
from pathlib import Path
from unittest.mock import MagicMock, patch

import docx as python_docx


# ─────────────────────────────────────────────
# Helpers: costruttori di DOCX minimi
# ─────────────────────────────────────────────

def _docx_vuoto(path: Path) -> None:
    """DOCX con solo un paragrafo vuoto (default python-docx)."""
    d = python_docx.Document()
    d.save(str(path))


def _docx_testo(path: Path, testo: str) -> None:
    """DOCX con un singolo paragrafo di testo."""
    d = python_docx.Document()
    d.add_paragraph(testo)
    d.save(str(path))


def _docx_tabella(path: Path) -> None:
    """DOCX con una tabella 2x2, nessun paragrafo di testo."""
    d = python_docx.Document()
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Nome"
    table.cell(0, 1).text = "Valore"
    table.cell(1, 0).text = "Progetto"
    table.cell(1, 1).text = "RAG VE"
    d.save(str(path))


def _docx_misto(path: Path) -> None:
    """DOCX con paragrafo di testo + tabella."""
    d = python_docx.Document()
    d.add_paragraph("Titolo documento di test")
    table = d.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Colonna A"
    table.cell(0, 1).text = "Colonna B"
    d.save(str(path))


# ─────────────────────────────────────────────
# Test 1: DOCX vuoto → stringa vuota, nessun crash
# ─────────────────────────────────────────────

def test_read_docx_file_vuoto_ritorna_stringa_vuota(tmp_path):
    """Un DOCX senza testo deve ritornare '' senza sollevare eccezioni."""
    p = tmp_path / "vuoto.docx"
    _docx_vuoto(p)

    from app.ingest_fs import read_docx_file

    result = read_docx_file(p)

    assert isinstance(result, str)
    assert result == ""


# ─────────────────────────────────────────────
# Test 2: DOCX con testo semplice → testo estratto correttamente
# ─────────────────────────────────────────────

def test_read_docx_file_testo_semplice_estratto(tmp_path):
    """Il testo del paragrafo deve comparire nella stringa restituita."""
    p = tmp_path / "testo.docx"
    _docx_testo(p, "Questo è un bando del Comune di Venezia.")

    from app.ingest_fs import read_docx_file

    result = read_docx_file(p)

    assert "bando" in result
    assert "Venezia" in result


# ─────────────────────────────────────────────
# Test 3: DOCX con tabella → celle estratte con separatore "|"
# ─────────────────────────────────────────────

def test_read_docx_file_tabella_estratta_con_separatore(tmp_path):
    """Le celle della tabella devono essere presenti nel testo con separatore '|'."""
    p = tmp_path / "tabella.docx"
    _docx_tabella(p)

    from app.ingest_fs import read_docx_file

    result = read_docx_file(p)

    assert "Nome" in result
    assert "Progetto" in result
    assert "RAG VE" in result
    assert "|" in result


# ─────────────────────────────────────────────
# Test 4: DOCX con paragrafo + tabella → entrambi estratti
# ─────────────────────────────────────────────

def test_read_docx_file_misto_paragrafo_e_tabella(tmp_path):
    """Sia i paragrafi sia le celle tabella devono essere presenti nel risultato."""
    p = tmp_path / "misto.docx"
    _docx_misto(p)

    from app.ingest_fs import read_docx_file

    result = read_docx_file(p)

    assert "Titolo documento di test" in result
    assert "Colonna A" in result
    assert "Colonna B" in result


# ─────────────────────────────────────────────
# Test 5: DOCX corrotto → eccezione sollevata, non inghiottita
# ─────────────────────────────────────────────

def test_read_docx_file_corrotto_solleva_eccezione(tmp_path):
    """Un file binario non valido deve sollevare un'eccezione (BadZipFile o simile).
    La funzione NON deve inghiottire silenziosamente l'errore.
    """
    p = tmp_path / "corrotto.docx"
    p.write_bytes(b"questo non e' un file DOCX valido: XYZABC123")

    from app.ingest_fs import read_docx_file

    with pytest.raises(Exception):
        read_docx_file(p)


# ─────────────────────────────────────────────
# Test 6: ingest_single_file con DOCX vuoto → status "skipped"
# ─────────────────────────────────────────────

def test_ingest_single_file_docx_vuoto_ritorna_skipped(tmp_path):
    """Se il DOCX è vuoto, ingest_single_file deve restituire status='skipped'
    senza inserire nulla nel DB.
    """
    p = tmp_path / "vuoto.docx"
    _docx_vuoto(p)

    # Mock minimale per get_conn: evita connessione DB reale.
    # ensure_kb: SELECT trova la KB esistente → ritorna kb_id.
    mock_cur = MagicMock()
    mock_cur.fetchone.return_value = ("kb-id-123",)

    mock_conn = MagicMock()
    mock_conn.cursor.return_value.__enter__ = MagicMock(return_value=mock_cur)
    mock_conn.cursor.return_value.__exit__ = MagicMock(return_value=False)
    mock_conn.autocommit = False

    with patch("app.ingest_fs.get_conn", return_value=mock_conn):
        from app.ingest_fs import ingest_single_file

        result = ingest_single_file(p, "demo")

    assert result["status"] == "skipped"
    assert "vuoto" in result.get("reason", "")
    # Nessun INSERT deve essere stato chiamato
    for call in mock_cur.execute.call_args_list:
        assert "INSERT INTO chunks" not in str(call), "Non devono esserci INSERT su DOCX vuoto"
